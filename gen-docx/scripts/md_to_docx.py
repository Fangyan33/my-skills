#!/usr/bin/env python3
"""Convert a constrained Markdown document into a styled DOCX file."""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Sequence

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


FONT_NAME = "仿宋"
BODY_SIZE = Pt(16)
TITLE_SIZE = Pt(22)
FOOTER_SIZE = Pt(9)
FIRST_LINE_INDENT_PT = Pt(32)

TABLE_HEADER_FONT = "黑体"
TABLE_BODY_FONT = "宋体"
TABLE_FONT_SIZE = Pt(10)
TABLE_HEADER_FILL = "D9EAF7"
TABLE_BORDER_COLOR = "999999"
TABLE_BORDER_SIZE = "4"
TABLE_CELL_MARGIN_TOP_BOTTOM = "100"
TABLE_CELL_MARGIN_LEFT_RIGHT = "120"

INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|==.+?==)")
IMAGE_ONLY_PATTERN = re.compile(r"^!\[(.*?)\]\((.+?)\)$")
IMAGE_ANYWHERE_PATTERN = re.compile(r"!\[(.*?)\]\((.+?)\)")
TABLE_SEPARATOR_PATTERN = re.compile(r"^\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?$")


@dataclass
class Block:
    kind: str
    text: str = ""
    marker: str = ""
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)
    image_path: str = ""
    alt_text: str = ""


@dataclass
class ParseResult:
    blocks: List[Block]
    warnings: List[str]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert a limited Markdown document to a styled DOCX file."
    )
    parser.add_argument("input", help="Path to the input Markdown file.")
    parser.add_argument("output", help="Path to the output DOCX file.")
    parser.add_argument("--title", help="Override the title read from Markdown.")
    parser.add_argument("--footer", help="Footer text to write into the document.")
    parser.add_argument("--author", help="Document author metadata.")
    parser.add_argument("--subject", help="Document subject metadata.")
    return parser.parse_args()


def read_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def split_table_cells(line: str) -> List[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]


def is_table_separator(line: str) -> bool:
    return bool(TABLE_SEPARATOR_PATTERN.match(line.strip()))


def parse_blocks(markdown_text: str) -> ParseResult:
    lines = markdown_text.splitlines()
    blocks: List[Block] = []
    warnings: List[str] = []
    paragraph_lines: List[str] = []
    in_code_block = False
    code_block_lines: List[str] = []
    code_block_start_line: Optional[int] = None
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        text = " ".join(line.strip() for line in paragraph_lines).strip()
        if text:
            blocks.append(Block("paragraph", text=text))
        paragraph_lines = []

    while index < len(lines):
        lineno = index + 1
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                code_text = " ".join(code_block_lines).strip()
                if code_text:
                    blocks.append(Block("paragraph", text=code_text))
                warnings.append(
                    f"Line {code_block_start_line}: code block downgraded to plain paragraph text."
                )
                in_code_block = False
                code_block_lines = []
                code_block_start_line = None
            else:
                flush_paragraph()
                in_code_block = True
                code_block_start_line = lineno
            index += 1
            continue

        if in_code_block:
            if stripped:
                code_block_lines.append(stripped)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if is_table_row(stripped):
            flush_paragraph()
            if index + 1 < len(lines) and is_table_separator(lines[index + 1].strip()):
                headers = split_table_cells(stripped)
                row_count = len(headers)
                table_rows: List[List[str]] = []
                table_end = index + 2
                malformed = False
                while table_end < len(lines):
                    table_line = lines[table_end].rstrip()
                    table_stripped = table_line.strip()
                    if not table_stripped:
                        break
                    if not is_table_row(table_stripped):
                        break
                    cells = split_table_cells(table_stripped)
                    if len(cells) != row_count:
                        malformed = True
                        break
                    table_rows.append(cells)
                    table_end += 1

                if malformed:
                    blocks.append(Block("paragraph", text=stripped))
                    warnings.append(
                        f"Line {lineno}: malformed GFM table downgraded to plain paragraph text."
                    )
                    index += 1
                    continue

                blocks.append(Block("table", headers=headers, rows=table_rows))
                index = table_end
                continue

            blocks.append(Block("paragraph", text=stripped))
            warnings.append(f"Line {lineno}: table row downgraded to plain paragraph text.")
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            blocks.append(Block("paragraph", text=stripped.lstrip(">").strip()))
            warnings.append(f"Line {lineno}: blockquote downgraded to plain paragraph text.")
            index += 1
            continue

        if re.match(r"^\s{2,}([-*]|\d+\.)\s+.+$", line):
            flush_paragraph()
            blocks.append(Block("paragraph", text=stripped))
            warnings.append(f"Line {lineno}: nested list downgraded to plain paragraph text.")
            index += 1
            continue

        image_match = IMAGE_ONLY_PATTERN.match(stripped)
        if image_match:
            flush_paragraph()
            blocks.append(
                Block("image", image_path=image_match.group(2).strip(), alt_text=image_match.group(1).strip())
            )
            index += 1
            continue

        if IMAGE_ANYWHERE_PATTERN.search(stripped):
            warnings.append(
                f"Line {lineno}: inline image mixed with text downgraded to plain paragraph text."
            )

        if stripped.startswith("# "):
            flush_paragraph()
            blocks.append(Block("title", text=stripped[2:].strip()))
            index += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            blocks.append(Block("heading1", text=stripped[3:].strip()))
            index += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            blocks.append(Block("heading2", text=stripped[4:].strip()))
            index += 1
            continue

        bullet_match = re.match(r"^([-*])\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            blocks.append(Block("list", text=bullet_match.group(2).strip(), marker="•"))
            index += 1
            continue

        ordered_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ordered_match:
            flush_paragraph()
            marker = f"{ordered_match.group(1)}、"
            blocks.append(Block("list", text=ordered_match.group(2).strip(), marker=marker))
            index += 1
            continue

        paragraph_lines.append(stripped)
        index += 1

    flush_paragraph()
    if in_code_block:
        code_text = " ".join(code_block_lines).strip()
        if code_text:
            blocks.append(Block("paragraph", text=code_text))
        warnings.append(
            f"Line {code_block_start_line}: unclosed code block downgraded to plain paragraph text."
        )
    return ParseResult(blocks=blocks, warnings=warnings)


def ensure_rfonts(run, font_name: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def set_run_font(
    run,
    size: Pt,
    *,
    font_name: str = FONT_NAME,
    bold: bool = False,
    highlight: bool = False,
) -> None:
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    ensure_rfonts(run, font_name)


def set_snap_to_grid(paragraph: Paragraph, enabled: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    snap = p_pr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        p_pr.append(snap)
    snap.set(qn("w:val"), "1" if enabled else "0")


def add_inline_runs(paragraph: Paragraph, text: str, *, size: Pt, default_bold: bool = False) -> None:
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, bold=True)
        elif part.startswith("==") and part.endswith("=="):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, bold=default_bold, highlight=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, bold=default_bold)


def add_plain_run(paragraph: Paragraph, text: str, *, font_name: str, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, TABLE_FONT_SIZE, font_name=font_name, bold=bold)


def configure_document(document: DocumentType) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def set_paragraph_box_style(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_snap_to_grid(paragraph, False)


def write_title(document: DocumentType, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_box_style(paragraph)
    add_inline_runs(paragraph, title, size=TITLE_SIZE, default_bold=True)


def write_heading(document: DocumentType, text: str, *, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_box_style(paragraph)
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT_PT if level == 1 else None
    add_inline_runs(paragraph, text, size=BODY_SIZE, default_bold=True)


def write_paragraph(document: DocumentType, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_box_style(paragraph)
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT_PT
    add_inline_runs(paragraph, text, size=BODY_SIZE)


def write_list_item(document: DocumentType, text: str, marker: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_box_style(paragraph)
    paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT_PT
    add_inline_runs(paragraph, f"{marker} {text}", size=BODY_SIZE)


def set_cell_fill(cell: _Cell, fill: Optional[str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if fill is None:
        if shd is not None:
            tc_pr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell: _Cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_element = tc_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_element)
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:color"), TABLE_BORDER_COLOR)
        edge_element.set(qn("w:sz"), "0")
        edge_element.set(qn("w:space"), "0")


def set_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = tbl_borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tbl_borders.append(edge_element)
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:color"), "auto")
        edge_element.set(qn("w:sz"), TABLE_BORDER_SIZE)
        edge_element.set(qn("w:space"), "0")


def set_table_cell_margins(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    margins = {
        "top": TABLE_CELL_MARGIN_TOP_BOTTOM,
        "bottom": TABLE_CELL_MARGIN_TOP_BOTTOM,
        "left": TABLE_CELL_MARGIN_LEFT_RIGHT,
        "right": TABLE_CELL_MARGIN_LEFT_RIGHT,
    }
    for edge, value in margins.items():
        edge_element = tbl_cell_mar.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            tbl_cell_mar.append(edge_element)
        edge_element.set(qn("w:w"), value)
        edge_element.set(qn("w:type"), "dxa")


def set_table_layout_autofit(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")


def apply_table_column_widths(table: Table, document: DocumentType, column_count: int) -> None:
    section = document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    width_per_column_emu = int(available_width / column_count)
    width_per_column_dxa = int(width_per_column_emu / 635)
    for row in table.rows:
        for cell in row.cells:
            cell.width = width_per_column_emu
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_per_column_dxa))
            tc_w.set(qn("w:type"), "dxa")


def configure_cell_paragraph(paragraph: Paragraph, *, centered: bool) -> None:
    set_paragraph_box_style(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT


def write_table_cell(cell: _Cell, text: str, *, is_header: bool) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    while paragraph.runs:
        paragraph._p.remove(paragraph.runs[0]._r)
    configure_cell_paragraph(paragraph, centered=is_header)
    font_name = TABLE_HEADER_FONT if is_header else TABLE_BODY_FONT
    add_plain_run(paragraph, text, font_name=font_name, bold=is_header)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_fill(cell, TABLE_HEADER_FILL if is_header else None)
    set_cell_borders(cell)


def write_table(document: DocumentType, block: Block) -> None:
    column_count = len(block.headers)
    table = document.add_table(rows=len(block.rows) + 1, cols=column_count)
    table.style = "Normal Table"
    table.autofit = True
    set_table_borders(table)
    set_table_cell_margins(table)
    set_table_layout_autofit(table)
    apply_table_column_widths(table, document, column_count)

    for col_idx, text in enumerate(block.headers):
        write_table_cell(table.cell(0, col_idx), text, is_header=True)
    for row_idx, row_values in enumerate(block.rows, start=1):
        for col_idx, text in enumerate(row_values):
            write_table_cell(table.cell(row_idx, col_idx), text, is_header=False)


def write_image(document: DocumentType, block: Block, *, base_dir: Path, warnings: List[str]) -> None:
    image_path = block.image_path
    if "://" in image_path or image_path.startswith("data:"):
        warnings.append(f"Image '{image_path}' is not a supported local path; downgraded to plain paragraph text.")
        write_paragraph(document, f"![{block.alt_text}]({image_path})")
        return

    resolved = (base_dir / image_path).resolve()
    if not resolved.exists() or not resolved.is_file():
        warnings.append(f"Image not found: {resolved}")
        write_paragraph(document, f"![{block.alt_text}]({image_path})")
        return

    section = document.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_box_style(paragraph)
    run = paragraph.add_run()
    try:
        shape = run.add_picture(str(resolved))
        if shape.width > max_width:
            shape.width = Emu(max_width)
    except Exception as exc:  # pragma: no cover - depends on image library support
        warnings.append(f"Failed to insert image '{resolved}': {exc}")
        paragraph._element.getparent().remove(paragraph._element)
        write_paragraph(document, f"![{block.alt_text}]({image_path})")


def set_footer(document: DocumentType, footer_text: str) -> None:
    for section in document.sections:
        paragraph = (
            section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        )
        if paragraph.runs:
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(footer_text)
        set_run_font(run, FOOTER_SIZE)


def apply_core_properties(document: DocumentType, *, author: Optional[str], subject: Optional[str]) -> None:
    if author:
        document.core_properties.author = author
    if subject:
        document.core_properties.subject = subject


def build_document(
    blocks: Sequence[Block],
    *,
    base_dir: Path,
    title_override: Optional[str] = None,
    footer: Optional[str] = None,
    author: Optional[str] = None,
    subject: Optional[str] = None,
) -> tuple[DocumentType, List[str]]:
    document = Document()
    render_warnings: List[str] = []
    configure_document(document)
    apply_core_properties(document, author=author, subject=subject)

    remaining_blocks = list(blocks)
    title = title_override
    if title is None and remaining_blocks and remaining_blocks[0].kind == "title":
        title = remaining_blocks.pop(0).text

    if title:
        write_title(document, title)

    for block in remaining_blocks:
        if block.kind == "title":
            write_heading(document, block.text, level=1)
        elif block.kind == "heading1":
            write_heading(document, block.text, level=1)
        elif block.kind == "heading2":
            write_heading(document, block.text, level=2)
        elif block.kind == "list":
            write_list_item(document, block.text, block.marker)
        elif block.kind == "table":
            write_table(document, block)
        elif block.kind == "image":
            write_image(document, block, base_dir=base_dir, warnings=render_warnings)
        else:
            write_paragraph(document, block.text)

    if footer:
        set_footer(document, footer)

    return document, render_warnings


def validate_paths(input_path: Path, output_path: Path) -> None:
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    if input_path.is_dir():
        raise IsADirectoryError(f"Input path is a directory: {input_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)


def emit_warnings(warnings: Sequence[str]) -> None:
    for warning in warnings:
        print(f"Warning: {warning}", file=sys.stderr)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()
    validate_paths(input_path, output_path)

    markdown_text = read_markdown(input_path)
    parse_result = parse_blocks(markdown_text)
    document, render_warnings = build_document(
        parse_result.blocks,
        base_dir=input_path.parent,
        title_override=args.title,
        footer=args.footer,
        author=args.author,
        subject=args.subject,
    )
    document.save(str(output_path))
    emit_warnings([*parse_result.warnings, *render_warnings])
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
